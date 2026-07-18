#!/usr/bin/env python3
"""Extract ordered plain text and source metadata from a .docx chapter."""
from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path
from docx import Document


def extract_docx(path: Path) -> str:
    doc = Document(path)
    # Preserve paragraph order and blank separators; do not edit the prose.
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--meta", type=Path)
    args = parser.parse_args()

    if not args.source.exists():
        parser.error(f"Source does not exist: {args.source}")
    if args.source.suffix.lower() != ".docx":
        parser.error("Only .docx input is currently supported")

    text = extract_docx(args.source)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(text, encoding="utf-8")

    if args.meta:
        metadata = {
            "filename": args.source.name,
            "format": "docx",
            "word_count": len(re.findall(r"\b[\w’'-]+\b", text)),
            "sha256": hashlib.sha256(args.source.read_bytes()).hexdigest(),
            "paragraph_count": len([p for p in Document(args.source).paragraphs if p.text.strip()]),
        }
        args.meta.parent.mkdir(parents=True, exist_ok=True)
        args.meta.write_text(json.dumps(metadata, indent=2), encoding="utf-8")

    print(f"Extracted {len(text)} characters to {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
